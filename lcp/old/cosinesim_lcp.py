import pandas as pd
import numpy as np
import networkx as nx
from pathlib import Path
import docx
import re
from transformers import AutoTokenizer, AutoModel
import torch
from sklearn.model_selection import train_test_split
from collections import defaultdict
from sklearn.metrics.pairwise import cosine_similarity

import matplotlib.pyplot as plt
import networkx as nx
import matplotlib.cm as cm




class HebrewLegalCitationPredictor:
    def __init__(self):
        print("Initializing predictor...")
        # Initialize AlephBERT
        self.tokenizer = AutoTokenizer.from_pretrained('onlplab/alephbert-base')
        self.mapping_file='/Users/liorb/Library/CloudStorage/OneDrive-post.bgu.ac.il/Thesis!!!/code/lcp/fe - MAPPING.csv'
        
        if torch.cuda.is_available():
            self.device = torch.device('cuda')
            print('cuda')
        else:
            self.device = torch.device('cpu')
            print('cpu')
            
        self.model = AutoModel.from_pretrained('onlplab/alephbert-base').to(self.device)
        print('alephbert-base')
        
        # Citation patterns
        self.citation_patterns = [
            r'ע"פ (\d+/\d+)',  # Criminal appeals
            r'ת"פ (\d+[-/]\d+[-/]\d+)',  # Criminal cases
            r'בג"ץ (\d+/\d+)',  # Supreme Court cases
            r'רע"פ (\d+/\d+)'  # Criminal permission requests
        ]
        
        # Initialize storage
        self.documents = {}
        self.citation_network = nx.DiGraph()
        self.document_embeddings = {}
    
    def clean_text(self, text):  
        """Clean Hebrew legal text"""
        # Remove direction markers
        text = re.sub(r'\{dir="[^"]+"\}', '', text)
        # Remove brackets
        text = re.sub(r'[\[\]]', '', text)
        # Remove extra whitespace
        text = ' '.join(text.split())
        return text.strip()
        
    def _extract_citations(self, file_path):
        """Extract citations from text using regex patterns"""
        data = {
            'extracted_paragraphs': [],
            'kept_paragraphs': [],
            'removed_paragraphs': [],
            'citations': []
        }
        
        cited_cases = []  # Initialize here
        cleaned_text = ""  # Initialize here
        
        try:
            # Read document
            doc = docx.Document(file_path)
            doc_id = Path(file_path).stem

            # First pass: Process paragraphs and citations
            all_citations = []
            full_text = []

            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                
                # Check for citations
                if 'ע"פ' in para_text or 'ת"פ' in para_text or 'עפ"ג' in para_text:
                    if len(para_text.strip()) > 0:
                        # Extract case citations as simple strings
                        para_citations = []  # Renamed from cited_cases to avoid confusion
                        
                        # Extract criminal appeals (ע"פ)
                        if 'ע"פ' in para_text:
                            matches = re.finditer(r'ע"פ (\d+/\d+)', para_text)
                            para_citations.extend(f'ע"פ {match.group(1)}' for match in matches)
                        
                        # Extract criminal cases (ת"פ)
                        if 'ת"פ' in para_text:
                            matches = re.finditer(r'ת"פ (\d+[-/]\d+[-/]\d+)', para_text)
                            para_citations.extend(f'ת"פ {match.group(1)}' for match in matches)
                        
                        # Extract עפ"ג cases
                        if 'עפ"ג' in para_text:
                            matches = re.finditer(r'עפ"ג (\d+/\d+)', para_text)
                            para_citations.extend(f'עפ"ג {match.group(1)}' for match in matches)
                        
                        # Store paragraph with simple citation list
                        data['extracted_paragraphs'].append({
                            'document': doc_id,
                            'paragraph_number': i,
                            'text': para_text,
                            'cited_cases': para_citations
                        })
                        
                        # Add to overall citations
                        all_citations.extend(para_citations)
                        cited_cases.extend(para_citations)  # Update the return value
                            
                else:
                    # Clean and store non-citation paragraphs
                    cleaned_text = self.clean_text(para_text)
                    if cleaned_text.strip():
                        data['kept_paragraphs'].append({
                            'document': doc_id,
                            'paragraph_number': i,
                            'text': cleaned_text
                        })
                        full_text.append(cleaned_text)

            cleaned_text = '\n'.join(full_text)  # Update the return value
            
            # Store document with simplified citation list
            self.documents[doc_id] = {
                'text': cleaned_text,
                'citations': list(set(all_citations)),  # Deduplicated list of citations
                'extracted_data': data
            }
        except Exception as e:
            print(f"Error processing document: {str(e)}")
        
        return list(set(cited_cases)), cleaned_text
    
            
    def _get_document_embedding(self, text):
        # Tokenize and get BERT embeddings
        inputs = self.tokenizer(
            text,
            padding=True,
            truncation=True,
            return_tensors="pt"
        )
        
        # Move inputs to device
        inputs = {k: v.to(self.device) for k, v in inputs.items()}
        
        outputs = self.model(**inputs)
        # Use [CLS] token representation
        return outputs.last_hidden_state[:, 0, :]


    def process_directory(self, docs_directory):
        """Process all documents in a directory"""
        print(f"Processing documents from: {docs_directory}")
        docs_path = Path(docs_directory)
        
        if not docs_path.exists():
            raise ValueError(f"Directory does not exist: {docs_directory}")
            
        # Process all docx files
        processed_docs = 0
        errors = []
        
        for file_path in docs_path.glob('*.docx'):
            try:
                result = self.process_document(file_path)
                if result:
                    processed_docs += 1
                    if processed_docs % 10 == 0:  # Progress update every 10 documents
                        print(f"Processed {processed_docs} documents...")
            except Exception as e:
                errors.append((file_path, str(e)))
                print(f"Error processing {file_path}: {str(e)}")
                
        print(f"Finished processing {processed_docs} documents")
        if errors:
            print("\nErrors occurred while processing the following files:")
            for file_path, error in errors:
                print(f"{file_path}: {error}")
        
        if processed_docs == 0:
            raise ValueError("No documents were processed. Check if directory contains .docx files")
            
        
    def process_document(self, file_path):
        """Process a single legal document"""
        try:
            # Ensure file_path is a Path object and exists
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            df = pd.read_csv(self.mapping_file)
            # Dropping unnecessary columns and cleaning the data
            df = df[['Mapping Value', 'Target Number']].dropna()

            # Creating the mapping_dict: Keys = Mapping Value, Values = Target Number
            mapping_dict = dict(zip(df['Mapping Value'], df['Target Number']))

    
            # Read document
            doc = docx.Document(file_path)
            doc_name = file_path.stem  # No extension
            doc_name_with_doc = doc_name + '.doc'  # With `.doc`
            doc_name_with_docx = file_path.name  # Full name with `.docx`

            # Check in the mapping_dict
            doc_id = (
                mapping_dict.get(doc_name) or
                mapping_dict.get(doc_name_with_doc) or
                mapping_dict.get(doc_name_with_docx) or
                doc_name  # Default to the original name if no match is found
)
            if doc_id==doc_name:
                print("doc_name fail to map:",doc_name)

            elif not doc_id.startswith('ת״פ '):
                doc_id = f'ת״פ {doc_id}'


            
            # Extract text and citations
            citations, text = self._extract_citations(file_path)
            
            # Store document only if we have valid data
            if text:  # Only store if we have text
                self.documents[doc_id] = {
                    'text': text,
                    'citations': citations
                }
                
                # Update citation network
                for citation in citations:
                    self.citation_network.add_edge(doc_id, citation)
                
            return doc_id
                
        except Exception as e:
            print(f"Error processing document {file_path}: {str(e)}")
            import traceback
            traceback.print_exc()  # This will print the full error trace
            return None    
        
    def build_model(self, docs_directory):
        """Build the complete model"""
        docs_path = Path(docs_directory)
        if not docs_path.exists():
            raise ValueError(f"Directory not found: {docs_directory}")
        
        print(f"Building model from documents in: {docs_directory}")
        # Process documents
        self.process_directory(docs_directory)
        
        print("Creating document embeddings...")
        # Create embeddings for documents
        for doc_id, doc_data in self.documents.items():
            self.document_embeddings[doc_id] = self._get_document_embedding(doc_data['text'])
            
        print("Calculating PageRank scores...")
        # Calculate PageRank scores
        self.pagerank_scores = nx.pagerank(self.citation_network)
        
        print("Model ready for predictions")

    def evaluate_model(self):
        """Run evaluation on test set"""
        results = {
            'precision': [],
            'recall': []
        }
        
        print("Evaluating model on test set...")
        for doc_id, doc_data in self.test_documents.items():
            # Get predictions
            predictions = set(self.predict_citations(doc_data['text']))
            actual = set(doc_data['citations'])
            
            # Calculate metrics
            if len(predictions) > 0:
                precision = len(predictions & actual) / len(predictions)
                results['precision'].append(precision)
            
            if len(actual) > 0:
                recall = len(predictions & actual) / len(actual)
                results['recall'].append(recall)
        
        metrics = {
            'precision': np.mean(results['precision']),
            'recall': np.mean(results['recall']),
            'f1': 2 * (np.mean(results['precision']) * np.mean(results['recall'])) / 
                  (np.mean(results['precision']) + np.mean(results['recall']))
        }
        
        print("\nTest Set Performance:")
        print(f"Precision: {metrics['precision']:.3f}")
        print(f"Recall: {metrics['recall']:.3f}")
        print(f"F1 Score: {metrics['f1']:.3f}")
        
        return metrics

    def predict_citations(self, text):
        """Predict citations for a given text"""
        text_embedding = self._get_document_embedding(text)
        similarities = {}
        for doc_id, doc_embedding in self.document_embeddings.items():
            e1 = text_embedding.detach().cpu().numpy().reshape(1, -1)
            e2 = doc_embedding.detach().cpu().numpy().reshape(1, -1)
            similarity = cosine_similarity(e1, e2)[0][0]
            similarities[doc_id] = similarity
        
        # Increase threshold to be more selective
        threshold = 0.95  
        
        # Take only top N most similar documents
        top_n = 5  # Limit number of similar documents considered
        
        predicted_citations = []
        sorted_docs = sorted(similarities.items(), key=lambda x: x[1], reverse=True)[:top_n]
        
        for doc_id, similarity in sorted_docs:
            if similarity > threshold:
                if doc_id in self.documents:
                    predicted_citations.extend(self.documents[doc_id]['citations'])
        
        return list(set(predicted_citations))
    
    # def display_entire_network(self):
    #     """Print all nodes and edges in the citation network."""
    #     print("\n--- Entire Citation Network ---")
    #     print(f"Total Nodes: {self.citation_network.number_of_nodes()}")
    #     print(f"Total Edges: {self.citation_network.number_of_edges()}")
        
    #     print("\nNodes:")
    #     for node in self.citation_network.nodes:
    #         print(node)
        
    #     print("\nEdges:")
    #     for edge in self.citation_network.edges:
    #         print(edge)

    # def display_test_set(self):
    #     """Print all documents in the test set."""
    #     print("\n--- Test Set ---")
    #     for doc_id, doc_data in self.test_documents.items():
    #         print(f"Document ID: {doc_id}")
    #         print(f"Text (excerpt): {doc_data['text'][:100]}...")  # Show only the first 100 characters
    #         print(f"Citations: {doc_data['citations']}")
    #         print("-" * 50)
    def save_entire_network(self, filename="entire_network.txt"):
        """Save all nodes and edges in the citation network to a file."""
        with open(filename, "w", encoding="utf-8") as file:
            file.write("--- Entire Citation Network ---\n")
            file.write(f"Total Nodes: {self.citation_network.number_of_nodes()}\n")
            file.write(f"Total Edges: {self.citation_network.number_of_edges()}\n\n")
            
            file.write("Nodes:\n")
            for node in self.citation_network.nodes:
                file.write(f"{node}\n")
            
            file.write("\nEdges:\n")
            for edge in self.citation_network.edges:
                file.write(f"{edge}\n")
        
        print(f"Entire network saved to {filename}")

    def save_test_set(self, filename="test_set.txt"):
        """Save the test set details to a file."""
        with open(filename, "w", encoding="utf-8") as file:
            file.write("--- Test Set ---\n")
            for doc_id, doc_data in self.test_documents.items():
                file.write(f"Document ID: {doc_id}\n")
                file.write(f"Text (excerpt): {doc_data['text'][:100]}...\n")  # Limit to 100 characters
                file.write(f"Citations: {doc_data['citations']}\n")
                file.write("-" * 50 + "\n")
        
        print(f"Test set details saved to {filename}")



    def visualize_citation_network(self, filename='citation_network.png', layout='spring', min_citations=0):
        """
        Visualize the citation network
        """
        if self.citation_network.number_of_nodes() == 0:
            print("No nodes in the citation network to visualize!")
            return
            
        # Create a copy of the graph for visualization
        viz_graph = self.citation_network.copy()
        
        # Calculate in-degrees for the original graph
        in_degrees = dict(viz_graph.in_degree())
        
        # Print debug information
        print("\nDebug - Before filtering:")
        print(f"Nodes: {viz_graph.number_of_nodes()}")
        print(f"Edges: {viz_graph.number_of_edges()}")
        print("Sample edges:", list(viz_graph.edges())[:5])
        
        # Remove nodes with fewer than min_citations incoming citations
        if min_citations > 0:
            nodes_to_remove = [node for node, degree in in_degrees.items() if degree < min_citations]
            viz_graph.remove_nodes_from(nodes_to_remove)
        
        # Print debug information after filtering
        print("\nDebug - After filtering:")
        print(f"Nodes: {viz_graph.number_of_nodes()}")
        print(f"Edges: {viz_graph.number_of_edges()}")
        print("Sample edges:", list(viz_graph.edges())[:5])
        
        if viz_graph.number_of_nodes() == 0:
            print(f"No nodes remain after filtering for minimum {min_citations} citations!")
            return
        
        # Create figure and axes
        fig, ax = plt.subplots(figsize=(20, 20))
        
        # Choose layout
        if layout == 'spring':
            pos = nx.spring_layout(viz_graph, k=1, iterations=50)
        else:
            pos = nx.circular_layout(viz_graph)
        
        # Calculate PageRank for node sizes
        pagerank = nx.pagerank(viz_graph)
        node_sizes = [pagerank[node] * 10000 for node in viz_graph.nodes()]
        
        # Calculate node colors based on in-degree
        current_in_degrees = dict(viz_graph.in_degree())
        max_degree = max(current_in_degrees.values()) if current_in_degrees else 0
        if max_degree > 0:
            node_colors = [current_in_degrees[node]/max_degree for node in viz_graph.nodes()]
        else:
            node_colors = [0 for node in viz_graph.nodes()]  # Default to 0 if no valid degrees
        
        # Draw the network
        nx.draw_networkx_nodes(viz_graph, pos, 
                            node_size=node_sizes,
                            node_color=node_colors, 
                            cmap=plt.cm.viridis,
                            ax=ax)
        
        nx.draw_networkx_edges(viz_graph, pos, 
                            edge_color='gray',
                            alpha=0.2,
                            arrows=True,
                            arrowsize=10,
                            ax=ax)
        
        # Add labels
        if pagerank:
            mean_pagerank = np.mean(list(pagerank.values()))
            labels = {node: node for node in viz_graph.nodes() 
                    if pagerank[node] > mean_pagerank}
            nx.draw_networkx_labels(viz_graph, pos, labels, font_size=8, ax=ax)
        
        # Add title
        ax.set_title(f'Citation Network (min {min_citations} citations)\n'
                    f'Nodes: {viz_graph.number_of_nodes()}, '
                    f'Edges: {viz_graph.number_of_edges()}')
        
        # Add colorbar
        sm = plt.cm.ScalarMappable(cmap=plt.cm.viridis)
        sm.set_array([])
        plt.colorbar(sm, ax=ax, label='Citation Count (Normalized)')
        
        # Remove axis
        ax.set_axis_off()
        
        # Save and close
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"\nGraph visualization saved as {filename}")
        
        # Print network statistics using the current in_degrees
        print("\nNetwork Statistics:")
        print(f"Total nodes (original): {self.citation_network.number_of_nodes()}")
        print(f"Total edges (original): {self.citation_network.number_of_edges()}")
        print(f"Nodes in visualization: {viz_graph.number_of_nodes()}")
        print(f"Edges in visualization: {viz_graph.number_of_edges()}")
        
        # Print top cited documents using the original in_degrees
        if in_degrees:
            top_cited = sorted(in_degrees.items(), key=lambda x: x[1], reverse=True)[:10]
            print("\nTop 10 most cited documents:")
            for doc, citations in top_cited:
                print(f"{doc}: {citations} citations")


    def split_acronym_and_number(self, verdict_id):
        """Split verdict ID into acronym and all remaining parts."""
        parts = verdict_id.split("-")
        if len(parts) > 1:
            acronym = parts[0]  # The first part (e.g., 'ME')
            number = " ".join(parts[1:])  # Join all remaining parts with a space
        else:
            acronym = "Unknown"  # Handle cases with no dash
            number = verdict_id
        return acronym, number
    
    def export_custom_document_citation_data(self, output_file="data.csv"):
        """Export document and citation data  to a CSV."""
        try:
            rows = []

            for doc_id, doc_data in self.documents.items():
                # Split document ID into acronym and verdict number
                doc_acronym, doc_verdict = self.split_acronym_and_number(doc_id)
                rows.append({
                        "Document Acronym": doc_acronym,
                        "Document Verdict": doc_verdict,
                    })
                
                # Process citations
                for citation in doc_data.get("citations", []):
                    citation_acronym, citation_verdict = self.split_acronym_and_number(citation)
                    rows.append({
                        "Citation Acronym": citation_acronym,
                        "Citation Verdict": citation_verdict
                    })

            # Convert to DataFrame
            df = pd.DataFrame(rows)

            # Save to CSV
            df.to_csv(output_file, index=False)
            print(f"Data exported to {output_file}")
        except Exception as e:
            print(f"Error exporting custom document and citation data: {e}")





if __name__ == '__main__':
    try:
        print("Starting the predictor...")
        predictor = HebrewLegalCitationPredictor()
        
        # Build the model with your documents
        predictor.build_model('/Users/liorb/Library/CloudStorage/GoogleDrive-liorkob@post.bgu.ac.il/.shortcut-targets-by-id/1f5AVMhCLkfM_ZoGYf_oDNiTxa2jVgL2B/חיזוי מתחמי ענישה/נתונים/New docx database/2018_cases_v1/docx')
        
        # Split documents into train and test sets
        doc_ids = list(predictor.documents.keys())
        train_ids, test_ids = train_test_split(doc_ids, test_size=0.2, random_state=42)
                # Add visualization with different options
        predictor.export_custom_document_citation_data()


        # Add this before calling visualize_citation_network
        print("\nDebug Citation Network Info:")
        print(f"Number of nodes: {predictor.citation_network.number_of_nodes()}")
        print(f"Number of edges: {predictor.citation_network.number_of_edges()}")
        print("Sample of nodes:", list(predictor.citation_network.nodes())[:5])
        print("Sample of edges:", list(predictor.citation_network.edges())[:5])

        predictor.visualize_citation_network(filename='citation_network_all.png', min_citations=1)
        predictor.visualize_citation_network(filename='citation_network_filtered.png', min_citations=3)
        predictor.visualize_citation_network(filename='citation_network_circular.png', 
                                          layout='circular', min_citations=2)

        # Create test_documents dictionary
        predictor.test_documents = {doc_id: predictor.documents[doc_id] 
                                  for doc_id in test_ids}
        
        # Run evaluation
        metrics = predictor.evaluate_model()

        # print("Display the entire network")
        # predictor.display_entire_network()

        # # Display the test set
        # predictor.display_test_set()
        predictor.save_entire_network("entire_network.txt")

        # Save the test set
        predictor.save_test_set("test_set.txt")

        
    except Exception as e:
        print(f"Error occurred: {str(e)}")
        import traceback
        traceback.print_exc()